import os
import shutil
import datetime

from django.conf import settings

import login.files_upload as files_upload
from docx import Document
import re


def create_job_folders(master_folder, job_number, job_title):
    job_folder_name = f"{job_number} {job_title}"
    job_folder_path = os.path.join(master_folder, job_folder_name)

    if os.path.exists(job_folder_path):
        shutil.rmtree(job_folder_path)
        print(f"Existing job folder '{job_folder_path}' removed.")

    os.makedirs(job_folder_path)

    subfolders = ['Quote', 'Visual', 'RAMS', 'Client Documents', 'PO']
    for subfolder in subfolders:
        subfolder_path = os.path.join(job_folder_path, subfolder)
        os.makedirs(subfolder_path)
        print(f"Created subfolder: {subfolder_path}")

    print(f"Created job folder: {job_folder_path}")


def replace_text_in_table(docx_path, replacements, job_number):
    doc = Document(docx_path)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for pattern, replacement in replacements.items():
                    new_text = re.sub(pattern, replacement, cell.text)
                    cell.text = new_text
    for section in doc.sections:
        footer = section.footer
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    if 'Texttexttext4' in cell.text:
                        cell.text = cell.text.replace('Texttexttext4', str(datetime.date.today()))
                    elif 'Texttexttext5' in cell.text:
                        cell.text = cell.text.replace('Texttexttext5', job_number)

    return doc


def main(jobNumber, jobDescription, address, dateOfWorks, duration, localHospital):
    master_folder = 'upload_files'
    job_number = jobNumber
    job_title = jobDescription
    folder_name = f"{job_number} {job_title}"
    create_job_folders(master_folder, job_number, job_title)
    document_path1 = os.path.join(settings.MEDIA_URL, 'document1.docx')
    document_path2 = os.path.join(settings.MEDIA_URL, 'document2.docx')

    input_docx_path1 = 'static/docs/document1.docx'
    output_docx_path1 = f'upload_files/{folder_name}/RAMS/{folder_name} - RA.docx'
    input_docx_path2 = 'static/docs/document2.docx'
    output_docx_path2 = f'upload_files/{folder_name}/RAMS/{folder_name} - MS.docx'

    job_number = jobNumber

    pattern_replacements = {
        r'\bTexttexttext1\b': folder_name,
        r'\bTexttexttext2\b': dateOfWorks,
        r'\bTexttexttext3\b': duration,
        r'\bTexttexttext4\b': str(datetime.date.today()),
        r'\bTexttexttext5\b': jobNumber,
        r'\bTexttexttext6\b': localHospital,
    }

    modified_docx = replace_text_in_table(input_docx_path1, pattern_replacements, job_number)
    modified_docx.save(output_docx_path1)
    modified_docx = replace_text_in_table(input_docx_path2, pattern_replacements, job_number)
    modified_docx.save(output_docx_path2)

    files_upload.main()
    return True
